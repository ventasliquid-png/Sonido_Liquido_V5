from docx import Document
import os

DOC_PATH = r"C:\dev\RAR_V1\Presupuesto Biotenk marzo 2025 a.docx"

def read_docx():
    if not os.path.exists(DOC_PATH):
        print(f"Error: File not found at {DOC_PATH}")
        return

    try:
        doc = Document(DOC_PATH)
        print(f"Document detected. Paragraphs: {len(doc.paragraphs)}")
        print("--- Content Preview ---")
        for p in doc.paragraphs[:5]:
            if p.text.strip():
                print(p.text)
        print("--- Tables ---")
        for i, table in enumerate(doc.tables):
            print(f"Table {i+1}: {len(table.rows)} rows, {len(table.columns)} cols")
            # Print first row content
            row_text = [cell.text.strip() for cell in table.rows[0].cells]
            print(f"  Header: {row_text}")
        print("-----------------------")
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    read_docx()
